import streamlit as st
import json
from datetime import datetime
import firebase_admin
from firebase_admin import credentials, firestore, storage
import uuid

# ==========================================
# 0. 初始化 Firebase (雲端資料庫)
# ==========================================
if not firebase_admin._apps:
    try:
        if "firebase" in st.secrets:
            firebase_secrets = dict(st.secrets["firebase"])
            cred = credentials.Certificate(firebase_secrets)
            
            # 取得 Storage Bucket 名稱並做自動去雜質處理
            storage_bucket_name = st.secrets.get("firebase_config", {}).get("storage_bucket", "")
            storage_bucket_name = storage_bucket_name.replace("gs://", "").strip()
            
            firebase_admin.initialize_app(cred, {
                'storageBucket': storage_bucket_name
            })
            st.warning("尚未設定 Firebase Secrets 金鑰。")
    except Exception as e:
        st.error(f"Firebase 初始化失敗：{e}")

try:
    db = firestore.client()
except:
    db = None

# ==========================================
# 引入 Google Calendar 模組
# ==========================================
try:
    from google_cal import add_diary_to_calendar
    GOOGLE_CAL_AVAILABLE = True
except ImportError:
    GOOGLE_CAL_AVAILABLE = False

# ==========================================
# 1. 基本設定
# ==========================================
st.set_page_config(page_title="DRKKY 的雲端日記", page_icon="☁️", layout="centered")

# ==========================================
# 2. 側邊欄：日期選擇
# ==========================================
st.sidebar.title("📅 日記導航")
selected_date = st.sidebar.date_input("選擇日期", datetime.today())
date_str_formatted = selected_date.strftime('%Y 年 %m 月 %d 日')
date_str_filename = selected_date.strftime('%Y-%m-%d')

# ==========================================
# 3. 讀取雲端資料
# ==========================================
existing_content = ""
existing_images_urls = []

if db:
    try:
        doc_ref = db.collection(u'diaries').document(date_str_filename)
        doc = doc_ref.get()
        
        if doc.exists:
            data = doc.to_dict()
            existing_content = data.get(u'content', "")
            existing_images_urls = data.get(u'image_urls', [])
    except Exception as e:
        st.error(f"讀取資料庫失敗：{e}")

st.title("☁️ DRKKY 的雲端日記")
st.subheader(f"📅 {date_str_formatted}")

# ==========================================
# 4. 編輯與顯示區域
# ==========================================
diary_text = st.text_area(
    label="今天過得如何？寫點什麼吧...",
    value=existing_content,
    height=350,
    placeholder="紀錄今天發生的趣事、心情或心得..."
)

if existing_images_urls:
    st.write("📸 **今日已儲存的圖片：**")
    cols = st.columns(3)
    for i, img_url in enumerate(existing_images_urls):
        cols[i % 3].image(img_url, use_container_width=True)
    
    if st.button("🗑️ 清除今日所有紀錄", key="clear_all_btn"):
        if db:
            doc_ref.delete()
            st.success("今日紀錄已從雲端清除！")
            st.rerun()

st.write("---")
uploaded_images = st.file_uploader("新增圖片 (將上傳至雲端)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

if uploaded_images:
    st.write("👀 **預覽即將上傳的圖片：**")
    cols_preview = st.columns(3)
    for i, img_file in enumerate(uploaded_images):
        cols_preview[i % 3].image(img_file, use_container_width=True)

# ==========================================
# 5. 儲存與同步邏輯
# ==========================================
st.divider()

col1, col2 = st.columns([1, 1])
with col1:
    save_btn = st.button("💾 儲存至雲端", type="primary", key="save_btn_main")
with col2:
    if GOOGLE_CAL_AVAILABLE:
        sync_to_google = st.checkbox("同時備份到 Google 行事曆", value=False)
    else:
        sync_to_google = False

if save_btn:
    if not db:
        st.error("無法連線至雲端資料庫，請檢查 Firebase 金鑰設定。")
    elif diary_text.strip() or uploaded_images or existing_images_urls:
        new_image_urls = []
        has_error = False  # 錯誤攔截器
        
        if uploaded_images:
            with st.spinner('正在上傳圖片至雲端儲存空間...'):
                try:
                    bucket = storage.bucket()
                    for img_file in uploaded_images:
                        file_extension = img_file.name.split('.')[-1]
                        unique_filename = f"diaries/{date_str_filename}_{uuid.uuid4().hex[:8]}.{file_extension}"
                        
                        blob = bucket.blob(unique_filename)
                        blob.upload_from_string(img_file.getvalue(), content_type=img_file.type)
                        blob.make_public()
                        new_image_urls.append(blob.public_url)
                except Exception as e:
                    st.error(f"圖片上傳失敗：{e}")
                    has_error = True 
        
        all_image_urls = existing_images_urls + new_image_urls
        
        if not has_error:
            with st.spinner('正在儲存日記內容...'):
                try:
                    doc_ref.set({
                        u'content': diary_text,
                        u'image_urls': all_image_urls,
                        u'updated_at': firestore.SERVER_TIMESTAMP
                    })
                    st.success("✅ 成功儲存至雲端資料庫！")
                except Exception as e:
                    st.error(f"儲存失敗：{e}")
                    has_error = True

        if sync_to_google and not has_error:
            with st.spinner('正在同步到 Google 行事曆...'):
                calendar_content = diary_text
                if all_image_urls:
                    calendar_content += f"\n\n[已附加 {len(all_image_urls)} 張圖片在雲端]"
                
                success, result = add_diary_to_calendar(date_str_filename, calendar_content)
                if success:
                    st.success(f"☁️ 行事曆同步成功！[點此查看]({result})")
                else:
                    st.error(f"❌ 行事曆同步失敗，錯誤訊息：{result}")
                    has_error = True
        
        if not has_error:
            st.rerun()
    else:
         st.warning("⚠️ 請先輸入內容或選擇圖片再儲存。")

# ==========================================
# 6. 實用小工具
# ==========================================
if diary_text:
    st.caption(f"✍️ 目前字數：{len(diary_text)} 字")
    
    st.write("---")
    st.write("📥 **匯出與下載**")
    
    col_ex1, col_ex2 = st.columns(2)
    
    txt_data = f"日期：{date_str_formatted}\n\n內容：\n{diary_text}"
    col_ex1.download_button(
        label="📄 下載純文字 (.txt)",
        data=txt_data,
        file_name=f"DRKKY_日記_{date_str_filename}.txt",
        mime="text/plain"
    )
    
    try:
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO
        import requests # 需要在 requirements.txt 加入 requests
        
        doc = Document()
        doc.add_heading('DRKKY 的雲端日記', 0)
        doc.add_heading(date_str_formatted, level=1)
        doc.add_paragraph(diary_text)
        
        # 修改後的圖片處理邏輯
        if existing_images_urls:
            doc.add_heading('附加圖片：', level=2)
            for url in existing_images_urls:
                try:
                    # 1. 從網址下載圖片內容
                    response = requests.get(url)
                    image_stream = BytesIO(response.content)
                    # 2. 將圖片插入 Word，寬度設定為 5 英吋
                    doc.add_picture(image_stream, width=Inches(5.0))
                except Exception as e:
                    doc.add_paragraph(f"無法載入圖片: {url}")
                
        bio = BytesIO()
        doc.save(bio)
        
        col_ex2.download_button(
            label="📝 下載 Word (.docx)",
            data=bio.getvalue(),
            file_name=f"DRKKY_日記_{date_str_filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except ImportError:
        col_ex2.warning("💡 請先在 requirements.txt 中加入 `python-docx` 才能啟用 Word 匯出功能。")
        
    st.info("💡 **如何將日記存為 PDF？**\n\n由於雲端環境的中文字體限制，最完美的 PDF 儲存方式是：直接按下電腦鍵盤的 `Ctrl + P` (Mac 為 `Cmd + P`)，將目的地選擇為 **「另存為 PDF」**，這樣能 100% 保留網頁上精美的深色排版與圖片喔！")